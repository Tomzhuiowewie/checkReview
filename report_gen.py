from pathlib import Path
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
import json,os
from typing import Dict, Any
import pandas as pd
from plot import plot_bar,plot_pie,plot_pic7
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches

def get_content_width_in_inches(docx_template_path):
    from docx import Document
    doc = Document(docx_template_path)
    section = doc.sections[0]
    width_emu = section.page_width - section.left_margin - section.right_margin
    return width_emu / 914400  # EMU 转英寸

def build_context(value, doc, template_path, content_width_inch=None, current_key=None):
    if content_width_inch is None:
        content_width_inch = get_content_width_in_inches(template_path)

    if isinstance(value, dict):
        return {
            k: build_context(v, doc, template_path, content_width_inch, current_key=k)
            for k, v in value.items()
        }

    elif isinstance(value, list):
        return [
            build_context(item, doc, template_path, content_width_inch, current_key=current_key)
            for item in value
        ]

    elif isinstance(value, str):
        ext = os.path.splitext(value)[1].lower()
        if ext in ['.jpg', '.jpeg', '.png'] and os.path.exists(value):
            # 仅非 pic3 和 pic7 才设置等宽
            if current_key not in ["pic3", "pic7", "pic4", "pic5", "pic6", "pic1", "pic2"]:
                return InlineImage(doc, value, width=Inches(content_width_inch))
            else:
                return InlineImage(doc, value)  # 原始尺寸
        return value

    else:
        return value

   

def insert_table_after_rendered_file(doc_path, table_data, placeholder="{{table_placeholder}}"):

    doc = Document(doc_path)
    found = False

    for i, paragraph in enumerate(doc.paragraphs):
        if placeholder in paragraph.text:
            found = True
            print(f"[Info] 找到占位符：{placeholder}，准备插入表格")
            # 删除该段落
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None

            # 插入表格标题
            if table_data.get("title"):
                doc.add_paragraph(table_data["title"], style="Heading 3")

            headers = table_data.get("header", [])
            rows = table_data.get("rows", [])

            if not headers or not isinstance(rows, list):
                raise ValueError("表格数据格式错误，缺少 header 或 rows")

            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"

            # 表头
            for col_idx, col in enumerate(headers):
                cell = table.cell(0, col_idx)
                cell.text = str(col)
                # 垂直居中
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # 水平居中
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 数据行
            for r_idx, row in enumerate(rows):
                # for c_idx, val in enumerate(row):
                #     table.cell(r_idx + 1, c_idx).text = str(val)
                for c_idx, val in enumerate(row):
                    cell = table.cell(r_idx + 1, c_idx)
                    cell.text = str(val)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break

    if not found:
        raise ValueError(f"未找到表格占位符：{placeholder}")

    doc.save(doc_path)

def generate_report(template_path, relate_data, pic_path, output_file_path):
    """
    template_path: 报告模版路径
    Config: 可配置的JSON数据
    pic_path: 图片保存位置
    out_file_path: 报告保存位置
    """
    # 1. 加载Word模板文件
    doc_tpl = DocxTemplate(template_path)

    for key, value in relate_data["pic_data"].items():
        save_path = pic_path + key + ".png"
        if key in ["pic1", "pic2", "pic5", "pic6"]:
            df = pd.DataFrame(value.items(), columns=['name', 'values'])
            plot_bar(df,save_path=save_path)
        elif key in ["pic3"]:
            df = pd.DataFrame(value.items(), columns=['name', 'values'])
            plot_pie(df,save_path=save_path)
        elif key in ["pic4"]:
            # 2. 生成报告所需图片  入围率（pie图信息）
            pic3 = pd.Series({k: v for k, v in relate_data["pic_data"]["pic3"].items() if isinstance(v, (int, float))})
            pic4 = pd.Series({k: v for k, v in relate_data["pic_data"]["pic4"].items() if isinstance(v, (int, float))})
            pic4_rate = (pic4 / pic3 * 100).round(2).to_dict()

            df = pd.DataFrame(value.items(), columns=['name', 'values'])
            plot_pie(df,save_path=save_path,paras=pic4_rate)
        elif key in ["pic7"]:
            df = pd.DataFrame(value.items(), columns=['name', 'values'])
            plot_pic7(df,save_path=save_path)

    # 3. 构建包含图片对象的 context
    # context = build_context(relate_data, doc_tpl)
    context = build_context(relate_data, doc_tpl, template_path)
    if "table_data" in relate_data:
        # 👇 添加这一行，避免占位符被 docxtpl 清除
        context["table_placeholder"] = "{{table_placeholder}}"
        # 5. 渲染模板
        doc_tpl.render(context)
        
        # 6. 保存生成的文件
        tmp_path = output_file_path.replace(".docx", "_filled.docx")

        doc_tpl.save(tmp_path)
        insert_table_after_rendered_file(tmp_path, relate_data["table_data"], placeholder="{{table_placeholder}}")
        # 4. 最终保存
        os.replace(tmp_path, output_file_path)
    else:
        doc_tpl.render(context)
        doc_tpl.save(output_file_path)

if __name__ == '__main__':
    # 可配置数据信息
    template_path = "./input/ReportTemplate/评审组组长AI分析报告json模版.json"
    Config_file = Path(template_path)
    with Config_file.open('r', encoding='utf-8') as f:
        raw_data = json.load(f)
    # 报告模版路径
    file_path = "./input/ReportTemplate/评审组组长AI分析报告.docx"
    # 输出报告路径
    output_file_path = "./output/report/评审组组长AI分析报告.docx"

    generate_report(file_path, raw_data, "./input/ReportTemplate/pic/", output_file_path)
